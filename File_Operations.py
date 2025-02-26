from __future__ import absolute_import, division, print_function, unicode_literals
from docx import Document
from dotenv import load_dotenv
load_dotenv()
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from pdfminer.pdfpage import PDFPage
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from typing import Dict, List, Union
import pypdf, pdfplumber, os
from pathlib import Path

from Config_Data import load_config_data

saved_dir = Path(os.getenv("FILE_PROCESSING_PATH"))
os.makedirs(saved_dir, exist_ok=True)

def iter_block_item(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def table_print(block):
    table_text = str()
    # table = block
    for row in block.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                table_text += f'{paragraph.text} \n'
    return table_text

def docx_para_table(file):
    text = str()
    document = Document(file)
    for block in iter_block_item(document):
        if isinstance(block, Paragraph):
            text += f'{block.text} \n'
        elif isinstance(block, Table):
            text += table_print(block)
    return text

def text_extract_pypdf(file: str) -> str:
    """
    Extracts text from a PDF file using PyPDF.

    Args:
        file (str): Path to the PDF file.

    Returns:
        str: Extracted text from the PDF or an error message if extraction fails.
    """
    try:
        text = ""
        pdf = pypdf.PdfReader(file)
        for page in pdf.pages:
            text += page.extract_text()
        return text
    except Exception as ep:
        return f"Something went wrong: {ep}"
    
def text_extract_pdfplumber(file: str) -> str:
    """
    Extracts text from a PDF file using pdfplumber.

    Args:
        file (str): Path to the PDF file.

    Returns:
        str: Extracted text from the PDF or an error message if extraction fails.
    """
    try:
        text = ""
        with pdfplumber.open(file) as pdf_pages:
            for page in pdf_pages.pages:
                text += page.extract_text(x_tolerance=1)
        return text
    except Exception as el:
        return f"Something went wrong: {el}"

def PDF_is_scanned(file: str) -> Union[bool, str]:
    """
    Determines whether a PDF file contains at least one scanned page.

    Args:
        file (str): Path to the PDF file.

    Returns:
        bool: True if the PDF has at least one scanned page, False otherwise.
        str: Error message if an exception occurs.
    """
    try:
        scanned_pages = []
        with open(file, 'rb') as infile:
            for page_num, page in enumerate(PDFPage.get_pages(infile), 1):
                if 'Font' not in page.resources.keys():
                    scanned_pages.append(page_num)
        return len(scanned_pages) > 0
    except Exception as eps:
        return f"Got the error: {eps}"
    
async def raw_text_from_file(file_path: str) -> str:
    """
    Extracts raw text from a given file, supporting DOCX and PDF formats.

    Args:
        file_path (str): The path to the input file.

    Returns:
        str: Extracted text from the file, or an error message if an issue occurs.
    """
    try:
        if os.path.basename(file_path).lower().endswith(".docx"):
            text = docx_para_table(file_path)
        elif os.path.basename(file_path).lower().endswith(".pdf"):
            if PDF_is_scanned(file_path):
                text = text_extract_pdfplumber(file_path)
            else:
                text = f"Provided pdf file {os.path.basename(file_path)} is a scanned file."
        else:
            raise ValueError(f"Unsupported file type {os.path.basename(file_path)}")
        return text
    except Exception as ex:
        return str(f"Error occurred {ex}")

def write_point_to_docx(data: Dict[str, List[str]], filename: str = "Output.docx") -> str:
    """
    Writes a dictionary to a DOCX file with headings and bullet points.

    Args:
        data (Dict[str, List[str]]): A dictionary where keys are section titles and values are lists of bullet points.
        filename (str, optional): The name of the output DOCX file. Defaults to "output.docx".

    Returns:
        str: Path of the generated PDF file.
    """
    complete_file_path = f"{saved_dir}\\{filename}"
    doc = Document()
    for title, points in data.items():
        doc.add_heading(title, level=1)
        for point in points:
            doc.add_paragraph(point, style='List Bullet')
    doc.save(complete_file_path)
    return f"Successfully wrote to {complete_file_path}"

def write_point_to_pdf(data: Dict[str, List[str]], filename: str = "Output.pdf") -> str:
    """
    Writes a dictionary to a PDF file with headings and bullet points.

    Args:
        data (Dict[str, List[str]]): A dictionary where keys are section titles and values are lists of bullet points.
        filename (str, optional): The name of the output PDF file. Defaults to "output.pdf".

    Returns:
        str: Path of the generated PDF file.
    """
    complete_file_path = f"{saved_dir}\\{filename}" 
    c = canvas.Canvas(complete_file_path, pagesize=letter)
    width, height = letter
    y_position = height - 50
    for title, points in data.items():
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y_position, title)
        y_position -= 20
        c.setFont("Helvetica", 12)
        for point in points:
            c.drawString(70, y_position, f"• {point}")
            y_position -= 20
        y_position -= 10
    c.save()
    return f"Successfully wrote to {complete_file_path}"